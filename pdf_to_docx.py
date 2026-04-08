# -*- coding: utf-8 -*-
"""将提取的文本转为格式化的 Word 文档"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 页边距
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

input_path = r'C:\Users\Administrator\.copaw\workspaces\CoPaw_QA_Agent_0.1beta1\nihaixia_shennong.txt'
output_path = r'C:\Users\Administrator\Downloads\倪海厦-人纪-神农百草经.docx'

with open(input_path, 'r', encoding='utf-8') as f:
    content = f.read()

# 按页分割
pages = re.split(r'\[Page \d+\]\s*\n?', content)
pages = [p.strip() for p in pages if p.strip()]

# 标题页 - 书籍名称
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('倪海厦 · 人纪')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('神农本草经')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

doc.add_paragraph('')
doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('（校勘版）\n\n')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

run2 = info.add_run('民间中医爱好者 敬校\n仅供学习交流，请勿用于商业用途')
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

for idx, page_text in enumerate(pages):
    if not page_text:
        continue
    
    page_num = idx + 1
    
    # 检查是否有类似标题的短行
    lines = page_text.split('\n')
    
    # 添加页码标记（方便对照原书）
    if any(lines):
        page_header = doc.add_paragraph()
        page_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = page_header.add_run(f'— {page_num} —')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    current_para = []
    for line in lines:
        line = line.strip()
        if not line:
            # 空行 = 段落结束
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
        elif re.match(r'^[一二三四五六七八九十百千零]+、', line) or \
             (len(line) < 30 and line not in ['']):
            # 可能是小节标题
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            h = doc.add_paragraph(line)
            # 不设为heading，保持正文样式但加粗
            for r in h.runs:
                r.bold = True
                r.font.size = Pt(14)
        else:
            current_para.append(line)
    
    if current_para:
        doc.add_paragraph(' '.join(current_para))

doc.save(output_path)
print(f'saved: {output_path}')
